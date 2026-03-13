import os
import shutil
import datetime
import db
from docx import Document
from docx.shared import Pt, Inches

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
REDE_OUTPUT_DIR = r"G:\NUMERADOR DADOS\OUTPUT"

MESES = {
    1: "01 - Janeiro", 2: "02 - Fevereiro", 3: "03 - Março",
    4: "04 - Abril", 5: "05 - Maio", 6: "06 - Junho",
    7: "07 - Julho", 8: "08 - Agosto", 9: "09 - Setembro",
    10: "10 - Outubro", 11: "11 - Novembro", 12: "12 - Dezembro"
}

TITULOS = {
    "OFICIO": "NUMERADOR DE OFÍCIO 2026",
    "MEMORANDO": "NUMERADOR DE MEMORANDO 2026",
    "CIRCULAR_INTERNA": "NUMERADOR DE CIRCULAR INTERNA 2026",
    "NOTIFICACAO": "NUMERADOR DE NOTIFICAÇÃO 2026",
    "PORTARIA": "NUMERADOR DE PORTARIA 2026",
    "AUTORIZACAO_VEICULO": "AUTORIZAÇÃO PARA CONDUÇÃO DE VEÍCULO OFICIAL 2026",
    "CERTIDAO": "NUMERADOR DE CERTIDÃO 2026"
}

NOME_UNICO = {
    "OFICIO": "Ofício",
    "MEMORANDO": "Memorando",
    "CIRCULAR_INTERNA": "Circular Interna",
    "NOTIFICACAO": "Notificação",
    "PORTARIA": "Portaria",
    "AUTORIZACAO_VEICULO": "Autorização de Veículo",
    "CERTIDAO": "Certidão"
}

def get_active_output_dir(tipo_db=None):
    base_out = REDE_OUTPUT_DIR if os.path.exists(r"G:\\") else OUTPUT_DIR
    
    hoje = datetime.datetime.now()
    ano = str(hoje.year)
    mes = MESES.get(hoje.month, "00 - Desconhecido")
    
    if tipo_db:
        nome_tipo = NOME_UNICO.get(tipo_db, tipo_db)
        caminho_final = os.path.join(base_out, "Docs Gerados", ano, mes, nome_tipo)
    else:
        caminho_final = os.path.join(base_out, "Docs Gerados", ano, mes)
        
    if not os.path.exists(caminho_final):
        try: os.makedirs(caminho_final)
        except: pass
        
    return caminho_final

def exportar_para_docx(tipo):
    target_dir = get_active_output_dir(tipo)
    file_path = os.path.join(target_dir, f'Relatorio_{NOME_UNICO.get(tipo, tipo)}_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    titulo_texto = TITULOS.get(tipo, f"NUMERADOR DE {tipo} 2026")
    titulo = doc.add_heading(str(), 0)
    run_titulo = titulo.add_run(titulo_texto)
    run_titulo.font.size = Pt(22)
    titulo.alignment = 1 
    
    registros = db.get_all_registros(tipo)
    registros = sorted(registros, key=lambda x: x[1])
    
    colunas = ['Nº', 'DATA', 'ASSUNTO', 'DESTINO', 'OBS', 'USUÁRIO']
    if tipo == "CERTIDAO":
        colunas.insert(1, 'PLACA')
        
    table = doc.add_table(rows=1, cols=len(colunas))
    table.style = 'Table Grid'
    table.allow_autofit = False
    
    hdr_cells = table.rows[0].cells
    for i, nome in enumerate(colunas):
        hdr_cells[i].text = nome
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                
    if tipo == "CERTIDAO":
        widths = [Inches(0.6), Inches(1.0), Inches(1.1), Inches(2.2), Inches(1.5), Inches(0.8), Inches(0.8)]
    else:
        widths = [Inches(0.6), Inches(1.1), Inches(2.8), Inches(1.5), Inches(1.0), Inches(1.0)]
        
    for idx, width in enumerate(widths):
        hdr_cells[idx].width = width
    
    for r in registros:
        p_id, numero, placa, data, assunto, destino, obs, usuario = r
        row_cells = table.add_row().cells
        
        row_cells[0].text = f"{numero:03d}"
        if tipo == "CERTIDAO":
            row_cells[1].text = placa or ""
            row_cells[2].text = data or ""
            row_cells[3].text = assunto or ""
            row_cells[4].text = destino or ""
            row_cells[5].text = obs or ""
            row_cells[6].text = usuario or ""
        else:
            row_cells[1].text = data or ""
            row_cells[2].text = assunto or ""
            row_cells[3].text = destino or ""
            row_cells[4].text = obs or ""
            row_cells[5].text = usuario or ""
            
        for i, width in enumerate(widths):
            row_cells[i].width = width
            
    doc.save(file_path)
    return file_path

