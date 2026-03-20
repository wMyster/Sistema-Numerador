from __future__ import annotations

import json
import os
import shutil
import sqlite3
import threading
from contextlib import contextmanager
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

import settings
from app_logger import get_logger

logger = get_logger(__name__)

BASE_DIR = settings.PROJECT_ROOT
NUMERADORES_DIR = os.path.join(BASE_DIR, "Numeradores")
REDE_DB_PATH = str(settings.get_network_data_dir() / "numerador.sqlite")


# ---------------------------------------------------------------------------
# Conexão
# ---------------------------------------------------------------------------

def get_active_db_path() -> str:
    return str(settings.get_db_path())


def connect() -> sqlite3.Connection:
    """Cria uma conexão com o banco SQLite com configurações padronizadas."""
    db_path = settings.get_db_path()
    db_dir = os.path.dirname(db_path)
    if not os.path.exists(db_dir):
        os.makedirs(db_dir)
    con = sqlite3.connect(str(db_path), timeout=10)
    con.row_factory = sqlite3.Row
    con.execute("PRAGMA foreign_keys = ON")
    return con


@contextmanager
def open_connection():
    """Fornece uma conexão SQLite que é sempre fechada ao fim do bloco."""
    con = connect()
    try:
        yield con
    finally:
        con.close()


def get_connection():
    """Compatibilidade retroativa — preferir open_connection() em código novo."""
    return connect()


# ---------------------------------------------------------------------------
# Migração de schema
# ---------------------------------------------------------------------------

def _ensure_columns(con: sqlite3.Connection, table: str, cols_sql: dict[str, str]) -> None:
    """Migração simples: se colunas não existirem, cria via ALTER TABLE."""
    existing = {row["name"] for row in con.execute(f"PRAGMA table_info({table})").fetchall()}
    for name, ddl in cols_sql.items():
        if name not in existing:
            con.execute(f"ALTER TABLE {table} ADD COLUMN {ddl}")


# ---------------------------------------------------------------------------
# Auditoria
# ---------------------------------------------------------------------------

def log_auditoria(
    usuario: str,
    acao: str,
    detalhes: str,
    *,
    actor: str | None = None,
    target_id: int | None = None,
) -> None:
    """Grava um evento de auditoria com informações enriquecidas."""
    agora = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    actor_final = (actor or usuario or "").strip() or None
    with open_connection() as con:
        con.execute("BEGIN IMMEDIATE")
        con.execute(
            """
            INSERT INTO auditoria (data_hora, usuario, acao, detalhes, actor, target_id)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (agora, usuario, acao, detalhes, actor_final, target_id),
        )
        con.commit()


def get_todas_auditorias(limit: int = 1000) -> list:
    with open_connection() as con:
        cur = con.execute(
            """
            SELECT data_hora, usuario, acao, detalhes
            FROM auditoria ORDER BY id DESC LIMIT ?
            """,
            (limit,),
        )
        return cur.fetchall()


# ---------------------------------------------------------------------------
# Schema
# ---------------------------------------------------------------------------

def init_schema(db_path: str) -> None:
    db_dir = os.path.dirname(db_path)
    if not os.path.exists(db_dir):
        os.makedirs(db_dir)

    with sqlite3.connect(db_path, timeout=10) as con:
        con.row_factory = sqlite3.Row
        con.execute("PRAGMA foreign_keys = ON")
        con.execute("BEGIN IMMEDIATE")

        # Tabela principal de registros
        con.execute(
            """
            CREATE TABLE IF NOT EXISTS registros (
                id         INTEGER PRIMARY KEY AUTOINCREMENT,
                tipo       TEXT    NOT NULL,
                numero     INTEGER NOT NULL,
                placa      TEXT,
                data       TEXT,
                assunto    TEXT,
                destino    TEXT,
                obs        TEXT,
                usuario    TEXT,
                created_at TEXT    NOT NULL DEFAULT (datetime('now','localtime')),
                updated_at TEXT,
                deleted_at TEXT,
                UNIQUE(tipo, numero)
            )
            """
        )

        # Garante colunas adicionadas em versões posteriores
        _ensure_columns(
            con,
            "registros",
            {
                "deleted_at": "deleted_at TEXT",
                "created_at": "created_at TEXT NOT NULL DEFAULT (datetime('now','localtime'))",
            },
        )

        # Tabela de auditoria enriquecida
        con.execute(
            """
            CREATE TABLE IF NOT EXISTS auditoria (
                id        INTEGER PRIMARY KEY AUTOINCREMENT,
                data_hora TEXT    NOT NULL,
                usuario   TEXT    NOT NULL,
                acao      TEXT    NOT NULL,
                detalhes  TEXT    NOT NULL,
                actor     TEXT,
                target_id INTEGER
            )
            """
        )

        _ensure_columns(
            con,
            "auditoria",
            {
                "actor": "actor TEXT",
                "target_id": "target_id INTEGER",
            },
        )

        # Índices de performance
        con.execute("CREATE INDEX IF NOT EXISTS idx_registros_tipo     ON registros(tipo)")
        con.execute("CREATE INDEX IF NOT EXISTS idx_registros_deleted  ON registros(deleted_at)")
        con.execute("CREATE INDEX IF NOT EXISTS idx_registros_data     ON registros(data)")
        con.execute("CREATE INDEX IF NOT EXISTS idx_registros_usuario  ON registros(usuario)")
        con.execute("CREATE INDEX IF NOT EXISTS idx_auditoria_data_hora ON auditoria(data_hora)")

        # Limpeza de registros sem dados úteis (migração legada)
        con.execute(
            """
            DELETE FROM registros
            WHERE deleted_at IS NULL
              AND (assunto IS NULL OR TRIM(assunto)  = '' OR TRIM(assunto)  = '-')
              AND (destino IS NULL OR TRIM(destino)  = '' OR TRIM(destino)  = '-')
              AND (obs     IS NULL OR TRIM(obs)      = '' OR TRIM(obs)      = '-')
              AND (usuario IS NULL OR TRIM(usuario)  = '' OR TRIM(usuario)  = '-')
            """
        )

        con.commit()


def init_db() -> None:
    logger.info("Inicializando banco de dados...")
    local_db = str(settings.LOCAL_DATA_DIR / "numerador.sqlite")
    init_schema(local_db)
    active_path = get_active_db_path()
    if active_path != local_db:
        init_schema(active_path)
        logger.info("Schema inicializado em local e rede")

    # Migração de users legados
    try:
        users_file = settings.get_users_file_path()
        if not users_file.exists():
            try:
                with open_connection() as con:
                    cur = con.execute("SELECT nome FROM usuarios")
                    nomes_existentes = [r["nome"].strip().upper() for r in cur.fetchall()]
                    users_dict = {u: "comum" for u in nomes_existentes if u}
                    users_dict.update({"DIRETORIA": "admin", "VIA DCT": "admin"})
                if not users_dict:
                    users_dict = {"DIRETORIA": "admin", "VIA DCT": "admin"}
            except Exception as e:
                logger.warning("Erro ao ler usuários do banco: %s", e)
                users_dict = {"DIRETORIA": "admin", "VIA DCT": "admin"}

            users_file.parent.mkdir(parents=True, exist_ok=True)
            with open(users_file, "w", encoding="utf-8") as f:
                json.dump(users_dict, f, indent=2, ensure_ascii=False)

            try:
                with open_connection() as con:
                    con.execute("BEGIN IMMEDIATE")
                    con.execute("DROP TABLE IF EXISTS usuarios")
                    con.commit()
            except Exception as e:
                logger.warning("Erro ao remover tabela usuarios legada: %s", e)
    except Exception as e:
        logger.error("Erro na migração de usuários legados: %s", e, exc_info=True)

    # Migração: move obs→usuario em registros antigos
    try:
        with open_connection() as con:
            con.execute("BEGIN IMMEDIATE")
            con.execute(
                """
                UPDATE registros
                SET usuario = UPPER(TRIM(obs)), obs = ''
                WHERE (usuario IS NULL OR TRIM(usuario) = '')
                  AND obs IS NOT NULL AND TRIM(obs) != '' AND LENGTH(obs) < 50
                """
            )
            con.commit()
    except Exception as e:
        logger.error("Erro na migração de obs->usuario: %s", e, exc_info=True)

    sincronizar_bancos()
    bootstrap_from_docx_if_empty()


# ---------------------------------------------------------------------------
# Backup e sincronização
# ---------------------------------------------------------------------------

def fazer_backup() -> None:
    try:
        from backup import perform_backup
        threading.Thread(target=perform_backup, kwargs={"is_manual": False}, daemon=True).start()
    except Exception as e:
        logger.error("Erro ao disparar backup em background: %s", e, exc_info=True)


def log_debug(msg: str) -> None:
    logger.debug(msg)


def sincronizar_bancos() -> None:
    rede_disponivel = settings.get_rede_path() is not None
    if rede_disponivel:
        rede_db = settings.get_network_data_dir() / "numerador.sqlite"
        local_db = settings.LOCAL_DATA_DIR / "numerador.sqlite"
        if local_db.exists() and rede_db.exists():
            merge_dbs(str(rede_db), str(local_db))
            merge_dbs(str(local_db), str(rede_db))
        elif rede_db.exists() and not local_db.exists():
            settings.LOCAL_DATA_DIR.mkdir(parents=True, exist_ok=True)
            shutil.copy2(str(rede_db), str(local_db))
            logger.info("Copiado DB da rede para o local.")
        elif local_db.exists() and not rede_db.exists():
            settings.get_network_data_dir().mkdir(parents=True, exist_ok=True)
            shutil.copy2(str(local_db), str(rede_db))
            logger.info("Copiado DB do local para a rede.")


def merge_dbs(db_main: str, db_attached: str) -> None:
    try:
        init_schema(db_main)
        init_schema(db_attached)
        with sqlite3.connect(db_main, timeout=10) as con:
            con.execute("BEGIN IMMEDIATE")
            con.execute(f"ATTACH DATABASE '{db_attached}' AS aux;")
            con.execute(
                """
                INSERT OR IGNORE INTO registros
                    (tipo, numero, placa, data, assunto, destino, obs, usuario, created_at, updated_at, deleted_at)
                SELECT tipo, numero, placa, data, assunto, destino, obs, usuario, created_at, updated_at, deleted_at
                FROM aux.registros
                WHERE (tipo, numero) NOT IN (SELECT tipo, numero FROM registros)
                """
            )
            con.execute(
                """
                UPDATE registros
                SET
                    placa      = (SELECT placa      FROM aux.registros a WHERE a.tipo = registros.tipo AND a.numero = registros.numero),
                    data       = (SELECT data       FROM aux.registros a WHERE a.tipo = registros.tipo AND a.numero = registros.numero),
                    assunto    = (SELECT assunto    FROM aux.registros a WHERE a.tipo = registros.tipo AND a.numero = registros.numero),
                    destino    = (SELECT destino    FROM aux.registros a WHERE a.tipo = registros.tipo AND a.numero = registros.numero),
                    obs        = (SELECT obs        FROM aux.registros a WHERE a.tipo = registros.tipo AND a.numero = registros.numero),
                    usuario    = (SELECT usuario    FROM aux.registros a WHERE a.tipo = registros.tipo AND a.numero = registros.numero),
                    updated_at = (SELECT updated_at FROM aux.registros a WHERE a.tipo = registros.tipo AND a.numero = registros.numero),
                    deleted_at = (SELECT deleted_at FROM aux.registros a WHERE a.tipo = registros.tipo AND a.numero = registros.numero)
                WHERE (tipo, numero) IN (
                    SELECT m.tipo, m.numero FROM registros m
                    JOIN aux.registros a ON m.tipo = a.tipo AND m.numero = a.numero
                    WHERE a.updated_at > m.updated_at
                )
                """
            )
            con.execute(
                """
                INSERT INTO auditoria (data_hora, usuario, acao, detalhes, actor, target_id)
                SELECT data_hora, usuario, acao, detalhes,
                       COALESCE(actor, usuario),
                       target_id
                FROM aux.auditoria
                WHERE (data_hora || acao || detalhes) NOT IN
                      (SELECT data_hora || acao || detalhes FROM auditoria)
                """
            )
            con.commit()
            logger.info("Bancos de dados mesclados: %s <- %s", db_main, db_attached)
    except Exception as e:
        logger.error("Erro ao mesclar bancos de dados %s e %s: %s", db_main, db_attached, e, exc_info=True)


# ---------------------------------------------------------------------------
# Bootstrap a partir de DOCX legado
# ---------------------------------------------------------------------------

def bootstrap_from_docx_if_empty() -> None:
    import docx as _docx

    mapa = {
        "NUMERADOR DE OFÍCIO 2026.docx": "OFICIO",
        "NUMERADOR DE MEMORANDO 2026.docx": "MEMORANDO",
        "NUMERADOR DE CIRCULAR INTERNA2026.docx": "CIRCULAR_INTERNA",
        "NUMERADOR DE NOTIFICAÇAO 2026.docx": "NOTIFICACAO",
        "NUMERADOR DE PORTARIA 2026.docx": "PORTARIA",
        "NUMERADOR DE AUTORIZAÇÃO PARA CONDUÇÃO DE VEÍCULO OFÍCIAL 2026.docx": "AUTORIZACAO_VEICULO",
        "NUMERADOR DE CERTIDÃO  2026.docx": "CERTIDAO",
    }

    with open_connection() as con:
        cur = con.execute("SELECT COUNT(*) AS cnt FROM registros WHERE deleted_at IS NULL")
        if cur.fetchone()["cnt"] > 0:
            logger.info("Banco de dados não está vazio, pulando bootstrap de DOCX.")
            return

    if not os.path.exists(NUMERADORES_DIR):
        logger.warning("Diretório de numeradores DOCX não encontrado: %s", NUMERADORES_DIR)
        return

    logger.info("Iniciando bootstrap de registros a partir de arquivos DOCX legados.")
    for nome_arquivo, tipo_db in mapa.items():
        filepath = os.path.join(NUMERADORES_DIR, nome_arquivo)
        if os.path.exists(filepath):
            try:
                doc = _docx.Document(filepath)
                if doc.tables:
                    table = doc.tables[0]
                    for row in table.rows[1:]:
                        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                        try:
                            n_str = cells[0].strip()
                            if not n_str:
                                continue
                            numero = int(n_str)
                        except Exception:
                            logger.warning("Não foi possível converter número de registro para int: %s", cells[0])
                            continue

                        placa = data = assunto = destino = obs = ""

                        if tipo_db == "CERTIDAO":
                            if len(cells) >= 6:
                                placa, data, assunto, destino, obs = (
                                    cells[1], cells[2], cells[3], cells[4], cells[5],
                                )
                        else:
                            if len(cells) >= 5:
                                data, assunto, destino, obs = (
                                    cells[1], cells[2], cells[3], cells[4],
                                )

                        try:
                            insert_registro(tipo_db, numero, placa, data, assunto, destino, obs, "", skip_sync=True)
                        except Exception as e:
                            logger.error("Erro inserindo registro %s - %d do DOCX %s: %s", tipo_db, numero, nome_arquivo, e)
            except Exception as e:
                logger.error("Erro ao ler DOCX %s: %s", nome_arquivo, e, exc_info=True)
        else:
            logger.debug("Arquivo DOCX não encontrado: %s", filepath)

    sincronizar_bancos()
    logger.info("Bootstrap de DOCX concluído.")


# ---------------------------------------------------------------------------
# Usuários
# ---------------------------------------------------------------------------

def get_all_usuarios() -> list[str]:
    users_file = settings.get_users_file_path()
    users = ["DIRETORIA", "VIA DCT"]
    if users_file.exists():
        try:
            with open(users_file, "r", encoding="utf-8") as f:
                data = json.load(f)
                users = list(data.keys())
        except Exception as e:
            logger.error("Erro ao carregar usuários do arquivo: %s", e)
    if not users:
        users = ["DIRETORIA"]
    return sorted(users)


def get_usuario_role(nome: str) -> str:
    users_file = settings.get_users_file_path()
    try:
        if users_file.exists():
            with open(users_file, "r", encoding="utf-8") as f:
                data = json.load(f)
                role = str(data.get(nome.upper(), "")).strip().lower()
                if role == "admin" or nome.upper() in ["DIRETORIA", "VIA DCT"]:
                    return "admin"
                return "comum"
    except Exception as e:
        logger.error("Erro ao obter role do usuário %s: %s", nome, e)
    return "admin" if nome.upper() in ["DIRETORIA", "VIA DCT"] else "comum"


def add_usuario(nome: str, role: str = "comum") -> None:
    users_file = settings.get_users_file_path()
    try:
        if users_file.exists():
            with open(users_file, "r", encoding="utf-8") as f:
                data = json.load(f)
        else:
            data = {"DIRETORIA": "admin", "VIA DCT": "admin"}
    except Exception as e:
        logger.error("Erro ao carregar dados de usuários para adicionar: %s", e)
        data = {}

    data[nome.upper()] = role
    users_file.parent.mkdir(parents=True, exist_ok=True)
    try:
        with open(users_file, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        logger.info("Usuário '%s' adicionado/atualizado com role '%s'.", nome, role)

        if settings.get_rede_path():
            try:
                loc = settings.LOCAL_DATA_DIR / "users.json"
                loc.parent.mkdir(parents=True, exist_ok=True)
                with open(loc, "w", encoding="utf-8") as f:
                    json.dump(data, f, indent=2, ensure_ascii=False)
                logger.debug("Usuários sincronizados para o local de rede.")
            except Exception as e:
                logger.error("Erro ao sincronizar arquivo de usuários para o local de rede: %s", e)
    except Exception as e:
        logger.error("Erro ao salvar arquivo de usuários: %s", e)


def delete_usuario(nome: str) -> None:
    users_file = settings.get_users_file_path()
    try:
        if users_file.exists():
            with open(users_file, "r", encoding="utf-8") as f:
                data = json.load(f)
        else:
            data = {}
        if nome.upper() in data:
            del data[nome.upper()]
            with open(users_file, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            logger.info("Usuário '%s' removido.", nome)
            if settings.get_rede_path():
                loc = settings.LOCAL_DATA_DIR / "users.json"
                if loc.exists():
                    with open(loc, "w", encoding="utf-8") as f:
                        json.dump(data, f, indent=2, ensure_ascii=False)
                    logger.debug("Usuários sincronizados para o local de rede após exclusão.")
        else:
            logger.warning("Tentativa de remover usuário inexistente: '%s'", nome)
    except Exception as e:
        logger.error("Erro ao remover usuário '%s': %s", nome, e)


# ---------------------------------------------------------------------------
# Registros — operações principais
# ---------------------------------------------------------------------------

def get_proximo_numero(tipo: str) -> int:
    with open_connection() as con:
        cur = con.execute(
            "SELECT MAX(numero) AS mx FROM registros WHERE tipo = ? AND deleted_at IS NULL",
            (tipo,),
        )
        row = cur.fetchone()
        return (row["mx"] + 1) if row["mx"] is not None else 1


def insert_registro(
    tipo: str,
    numero: int,
    placa: str,
    data: str,
    assunto: str,
    destino: str,
    obs: str,
    usuario: str,
    skip_sync: bool = False,
) -> None:
    now = datetime.now().isoformat()
    with open_connection() as con:
        con.execute("BEGIN IMMEDIATE")
        con.execute(
            """
            INSERT OR IGNORE INTO registros
                (tipo, numero, placa, data, assunto, destino, obs, usuario, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (tipo, numero, placa, data, assunto, destino, obs, usuario, now, now),
        )
        con.commit()

    logger.info("Registro inserido: %s #%d", tipo, numero)
    if not skip_sync:
        sincronizar_bancos()
        fazer_backup()


def update_registro(
    id_registro: int,
    placa: str,
    data: str,
    assunto: str,
    destino: str,
    obs: str,
    usuario: str,
) -> None:
    now = datetime.now().isoformat()
    with open_connection() as con:
        con.execute("BEGIN IMMEDIATE")
        con.execute(
            """
            UPDATE registros
            SET placa = ?, data = ?, assunto = ?, destino = ?, obs = ?, usuario = ?, updated_at = ?
            WHERE id = ?
            """,
            (placa, data, assunto, destino, obs, usuario, now, id_registro),
        )
        con.commit()
    logger.info("Registro atualizado (ID: %d)", id_registro)
    sincronizar_bancos()
    fazer_backup()


def delete_registro(id_registro: int) -> None:
    """Soft-delete: marca deleted_at em vez de excluir fisicamente."""
    now = datetime.now().isoformat()
    with open_connection() as con:
        con.execute("BEGIN IMMEDIATE")
        con.execute(
            "UPDATE registros SET deleted_at = ?, updated_at = ? WHERE id = ?",
            (now, now, id_registro),
        )
        con.commit()
    logger.info("Registro enviado para lixeira (ID: %d)", id_registro)
    sincronizar_bancos()
    fazer_backup()


def restore_registro(id_registro: int) -> None:
    """Restaura um registro da lixeira."""
    now = datetime.now().isoformat()
    with open_connection() as con:
        con.execute("BEGIN IMMEDIATE")
        con.execute(
            "UPDATE registros SET deleted_at = NULL, updated_at = ? WHERE id = ?",
            (now, id_registro),
        )
        con.commit()
    logger.info("Registro restaurado da lixeira (ID: %d)", id_registro)
    sincronizar_bancos()
    fazer_backup()


def delete_permanente(id_registro: int) -> None:
    """Exclui fisicamente um registro (apenas da lixeira)."""
    with open_connection() as con:
        con.execute("BEGIN IMMEDIATE")
        con.execute("DELETE FROM registros WHERE id = ?", (id_registro,))
        con.commit()
    logger.info("Registro excluído permanentemente (ID: %d)", id_registro)
    sincronizar_bancos()
    fazer_backup()


def esvaziar_lixeira() -> int:
    """Exclui permanentemente todos os registros na lixeira. Retorna a quantidade."""
    with open_connection() as con:
        con.execute("BEGIN IMMEDIATE")
        cur = con.execute("DELETE FROM registros WHERE deleted_at IS NOT NULL")
        count = int(cur.rowcount or 0)
        con.commit()
    logger.info("%d registros excluídos permanentemente da lixeira.", count)
    sincronizar_bancos()
    fazer_backup()
    return count


# ---------------------------------------------------------------------------
# Consultas
# ---------------------------------------------------------------------------

def get_all_registros(
    tipo: str,
    busca: str = "",
    data_inicio: str = "",
    data_fim: str = "",
    include_deleted: bool = False,
    limit: int = 200,
) -> list:
    with open_connection() as con:
        query = (
            "SELECT id, numero, placa, data, assunto, destino, obs, usuario "
            "FROM registros WHERE tipo = ?"
        )
        params: list[Any] = [tipo]

        if not include_deleted:
            query += " AND deleted_at IS NULL"

        if busca:
            termo = f"%{busca}%"
            query += (
                " AND (assunto LIKE ? OR destino LIKE ? OR obs LIKE ? OR usuario LIKE ? OR placa LIKE ?)"
            )
            params.extend([termo, termo, termo, termo, termo])

        if data_inicio and data_fim:
            query += (
                " AND (SUBSTR(data,7,4)||'-'||SUBSTR(data,4,2)||'-'||SUBSTR(data,1,2)) BETWEEN ? AND ?"
            )
            try:
                di = datetime.strptime(data_inicio, "%d/%m/%Y").strftime("%Y-%m-%d")
                df = datetime.strptime(data_fim, "%d/%m/%Y").strftime("%Y-%m-%d")
                params.extend([di, df])
            except Exception as e:
                logger.warning("Erro ao parsear datas de busca '%s' e '%s': %s. Ignorando filtro de data.", data_inicio, data_fim, e)
                query = query.rsplit(" AND ", 1)[0]

        query += " ORDER BY numero DESC"
        if limit and limit > 0:
            query += " LIMIT ?"
            params.append(limit)

        cur = con.execute(query, tuple(params))
        return cur.fetchall()


def list_lixeira(tipo: str | None = None, limit: int = 200) -> list:
    """Lista registros apagados (soft-delete), opcionalmente filtrados por tipo."""
    with open_connection() as con:
        query = (
            "SELECT id, tipo, numero, placa, data, assunto, destino, obs, usuario, deleted_at "
            "FROM registros WHERE deleted_at IS NOT NULL"
        )
        params: list[Any] = []
        if tipo:
            query += " AND tipo = ?"
            params.append(tipo)
        query += " ORDER BY deleted_at DESC, id DESC"
        
        if limit and limit > 0:
            query += " LIMIT ?"
            params.append(limit)
            
        return con.execute(query, tuple(params)).fetchall()


def get_estatisticas() -> list:
    with open_connection() as con:
        cur = con.execute(
            """
            SELECT tipo, COUNT(*) AS cnt
            FROM registros
            WHERE deleted_at IS NULL
              AND (
                  (assunto IS NOT NULL AND TRIM(assunto) != '')
               OR (destino IS NOT NULL AND TRIM(destino) != '')
               OR (obs     IS NOT NULL AND TRIM(obs)     != '')
               OR (usuario IS NOT NULL AND TRIM(usuario) != '')
              )
            GROUP BY tipo
            ORDER BY tipo ASC
            """
        )
        return cur.fetchall()


# ---------------------------------------------------------------------------
# Analytics / Dashboard
# ---------------------------------------------------------------------------

def get_estatisticas_dashboard() -> dict:
    """Retorna estatísticas gerenciais do mês corrente."""
    hoje = datetime.now()
    mes_atual = f"{hoje.month:02d}/{hoje.year}"
    stats: dict = {"total_mes": 0, "por_tipo": {}, "usuarios_ativos": 0}

    try:
        with open_connection() as con:
            cur = con.execute(
                """
                SELECT tipo, COUNT(id) AS cnt FROM registros
                WHERE data LIKE ? AND deleted_at IS NULL
                GROUP BY tipo
                """,
                (f"%/{mes_atual}",),
            )
            for row in cur.fetchall():
                stats["por_tipo"][row["tipo"]] = row["cnt"]
                stats["total_mes"] += row["cnt"]

            cur = con.execute(
                """
                SELECT COUNT(DISTINCT usuario) AS cnt FROM registros
                WHERE data LIKE ? AND deleted_at IS NULL
                  AND usuario IS NOT NULL AND usuario != ''
                """,
                (f"%/{mes_atual}",),
            )
            row = cur.fetchone()
            if row:
                stats["usuarios_ativos"] = row["cnt"]
    except Exception as e:
        logger.error("Erro ao buscar estatísticas: %s", e, exc_info=True)

    return stats


def get_top_destinos(limite: int = 5) -> list:
    """Retorna os destinos mais frequentes."""
    try:
        with open_connection() as con:
            cur = con.execute(
                """
                SELECT destino, COUNT(id) AS qtd FROM registros
                WHERE deleted_at IS NULL
                  AND destino IS NOT NULL AND TRIM(destino) != ''
                GROUP BY TRIM(LOWER(destino))
                ORDER BY qtd DESC
                LIMIT ?
                """,
                (limite,),
            )
            return [(str(r["destino"]).strip().title(), r["qtd"]) for r in cur.fetchall()]
    except Exception as e:
        logger.error("Erro ao buscar top destinos: %s", e, exc_info=True)
    return []


def get_historico_assuntos(limite: int = 30) -> list[str]:
    try:
        with open_connection() as con:
            cur = con.execute(
                """
                SELECT DISTINCT TRIM(assunto) AS a FROM registros
                WHERE deleted_at IS NULL
                  AND assunto IS NOT NULL AND TRIM(assunto) != ''
                ORDER BY id DESC LIMIT ?
                """,
                (limite,),
            )
            return [r["a"] for r in cur.fetchall()]
    except Exception as e:
        logger.error("Erro ao buscar histórico de assuntos: %s", e)
    return []


def get_historico_destinos(limite: int = 30) -> list[str]:
    try:
        with open_connection() as con:
            cur = con.execute(
                """
                SELECT DISTINCT TRIM(destino) AS d FROM registros
                WHERE deleted_at IS NULL
                  AND destino IS NOT NULL AND TRIM(destino) != ''
                ORDER BY id DESC LIMIT ?
                """,
                (limite,),
            )
            return [r["d"] for r in cur.fetchall()]
    except Exception as e:
        logger.error("Erro ao buscar histórico de destinos: %s", e)
    return []


# ---------------------------------------------------------------------------
# Modelos Favoritos
# ---------------------------------------------------------------------------

def get_modelos_file_path() -> Path:
    p = settings.get_users_file_path().parent / "modelos_favoritos.json"
    if not p.exists():
        try:
            p.parent.mkdir(parents=True, exist_ok=True)
            with open(p, "w", encoding="utf-8") as f:
                json.dump({}, f)
            logger.info("Arquivo de modelos favoritos criado: %s", p)
        except Exception as e:
            logger.error("Erro ao criar arquivo de modelos favoritos: %s", e)
    return p


def get_modelos(tipo_db: str) -> dict:
    try:
        arquivo = get_modelos_file_path()
        if arquivo.exists():
            with open(arquivo, "r", encoding="utf-8") as f:
                data = json.load(f)
                return data.get(tipo_db, {})
    except Exception as e:
        logger.error("Erro ao carregar modelos favoritos para tipo '%s': %s", tipo_db, e)
    return {}


def save_modelo(tipo_db: str, nome_modelo: str, assunto: str, destino: str, obs: str) -> None:
    arquivo = get_modelos_file_path()
    try:
        if arquivo.exists():
            with open(arquivo, "r", encoding="utf-8") as f:
                data = json.load(f)
        else:
            data = {}
    except Exception as e:
        logger.error("Erro ao carregar dados de modelos favoritos para salvar: %s", e)
        data = {}

    if tipo_db not in data:
        data[tipo_db] = {}
    data[tipo_db][nome_modelo] = {"assunto": assunto, "destino": destino, "obs": obs}

    try:
        arquivo.parent.mkdir(parents=True, exist_ok=True)
        with open(arquivo, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        logger.info("Modelo favorito '%s' salvo para tipo '%s'.", nome_modelo, tipo_db)
        if settings.get_rede_path():
            loc = settings.LOCAL_DATA_DIR / "modelos_favoritos.json"
            loc.parent.mkdir(parents=True, exist_ok=True)
            with open(loc, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            logger.debug("Modelos favoritos sincronizados para o local de rede.")
    except Exception as e:
        logger.error("Erro ao salvar modelo favorito: %s", e, exc_info=True)


def delete_modelo(tipo_db: str, nome_modelo: str) -> None:
    arquivo = get_modelos_file_path()
    try:
        if arquivo.exists():
            with open(arquivo, "r", encoding="utf-8") as f:
                data = json.load(f)
            if tipo_db in data and nome_modelo in data[tipo_db]:
                del data[tipo_db][nome_modelo]
                with open(arquivo, "w", encoding="utf-8") as f:
                    json.dump(data, f, indent=2, ensure_ascii=False)
                if settings.get_rede_path():
                    loc = settings.LOCAL_DATA_DIR / "modelos_favoritos.json"
                    loc.parent.mkdir(parents=True, exist_ok=True)
                    with open(loc, "w", encoding="utf-8") as f:
                        json.dump(data, f, indent=2, ensure_ascii=False)
    except Exception:
        pass
